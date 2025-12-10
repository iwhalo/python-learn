"""
斗鱼文档内容抓取工具
用于抓取 https://doc.douyu.tv/ddse/preview/space/251381?sid=539 页面内容
"""
import requests
import json
import time
import os
from urllib.parse import urlencode
import re
import markdown
from docx import Document
import getpass

# 尝试导入weasyprint，如果失败则PDF功能不可用
WEASYPRINT_AVAILABLE = False
HTML = None
CSS = None

try:
    import weasyprint
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except OSError:
    # 处理库加载错误，如Windows上的DLL问题
    WEASYPRINT_AVAILABLE = False
    HTML = None
    CSS = None
except ImportError:
    # 处理导入错误
    WEASYPRINT_AVAILABLE = False
    HTML = None
    CSS = None


class DouyuDocScraper:
    def __init__(self):
        self.base_url = "https://doc.douyu.tv"
        self.session = requests.Session()
        # 设置通用请求头
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/142.0.0.0 Safari/537.36',
            'Accept': '*/*',
            'Accept-Language': 'zh-CN,zh;q=0.9,en-GB;q=0.8,en;q=0.7,en-US;q=0.6',
            'Accept-Encoding': 'gzip, deflate, br, zstd',
            'Referer': 'https://doc.douyu.tv/ddse/preview/space/251381?sid=539',
            'Sec-Ch-Ua': '"Chromium";v="142", "Google Chrome";v="142", "Not_A Brand";v="99"',
            'Sec-Ch-Ua-Mobile': '?0',
            'Sec-Ch-Ua-Platform': '"Windows"',
            'Sec-Fetch-Dest': 'empty',
            'Sec-Fetch-Mode': 'cors',
            'Sec-Fetch-Site': 'same-origin',
            'Cache-Control': 'no-cache',
            'Pragma': 'no-cache'
        })
        
        # 初始化登录状态
        self.is_logged_in = False
        self.login_cookies = {}
        
        # 从配置文件加载cookie
        self.load_cookies_from_config()
    
    def get_doc_content(self, res_id, sid, is_view=1):
        """
        获取文档内容
        :param res_id: 资源ID
        :param sid: 空间ID
        :param is_view: 是否为预览模式
        :return: 文档数据
        """
        url = f"{self.base_url}/v1/api/doc/getDoc"
        params = {
            'isView': is_view,
            'r': str(time.time()),  # 随机数参数
            'resId': res_id,
            'sid': sid
        }
        
        try:
            response = self.session.get(url, params=params)
            response.raise_for_status()
            data = response.json()
            
            if data.get('error') == 0:
                return data.get('data')
            else:
                error_msg = data.get('msg', '未知错误')
                print(f"获取文档内容失败: {error_msg}")
                
                # 检查是否是登录相关错误
                if '登录' in str(error_msg) or 'login' in str(error_msg).lower():
                    self.handle_login_required_error(error_msg)
                
                return None
        except requests.RequestException as e:
            print(f"请求文档内容时发生错误: {e}")
            return None
    
    def get_user_info(self):
        """
        获取用户信息
        :return: 用户数据
        """
        url = f"{self.base_url}/v1/api/user/getUserInfo"
        
        try:
            response = self.session.get(url)
            response.raise_for_status()
            data = response.json()
            
            if data.get('error') == 0:
                return data.get('data')
            else:
                error_msg = data.get('msg', '未知错误')
                print(f"获取用户信息失败: {error_msg}")
                
                # 检查是否是登录相关错误
                if '登录' in str(error_msg) or 'login' in str(error_msg).lower():
                    self.handle_login_required_error(error_msg)
                
                return None
        except requests.RequestException as e:
            print(f"请求用户信息时发生错误: {e}")
            return None
    
    def get_attachments(self, res_id, sid):
        """
        获取文档附件
        :param res_id: 资源ID
        :param sid: 空间ID
        :return: 附件列表
        """
        url = f"{self.base_url}/v1/api/doc/getAttachment"
        params = {
            'resId': res_id,
            'sid': sid
        }
        
        try:
            response = self.session.get(url, params=params)
            response.raise_for_status()
            data = response.json()
            
            if data.get('error') == 0:
                return data.get('data', [])
            else:
                error_msg = data.get('msg', '未知错误')
                print(f"获取附件失败: {error_msg}")
                
                # 检查是否是登录相关错误
                if '登录' in str(error_msg) or 'login' in str(error_msg).lower():
                    self.handle_login_required_error(error_msg)
                
                return []
        except requests.RequestException as e:
            print(f"请求附件时发生错误: {e}")
            return []
    
    def check_collect_status(self, doc_id, sid):
        """
        检查文档收藏状态
        :param doc_id: 文档ID
        :param sid: 空间ID
        :return: 收藏状态
        """
        url = f"{self.base_url}/v1/api/square/checkCollect"
        params = {
            'docId': doc_id,
            'sid': sid,
            'sourceType': 0
        }
        
        try:
            response = self.session.get(url, params=params)
            response.raise_for_status()
            data = response.json()
            
            if data.get('error') == 0:
                return data.get('data', {}).get('isCollect', 0)
            else:
                error_msg = data.get('msg', '未知错误')
                print(f"检查收藏状态失败: {error_msg}")
                
                # 检查是否是登录相关错误
                if '登录' in str(error_msg) or 'login' in str(error_msg).lower():
                    self.handle_login_required_error(error_msg)
                
                return 0
        except requests.RequestException as e:
            print(f"请求收藏状态时发生错误: {e}")
            return 0
    
    def check_subscribe_status(self, res_id):
        """
        检查订阅状态
        :param res_id: 资源ID
        :return: 订阅状态
        """
        url = f"{self.base_url}/v1/api/doc/checkSubscribe"
        params = {
            'resId': res_id
        }
        
        try:
            response = self.session.get(url, params=params)
            response.raise_for_status()
            data = response.json()
            
            if data.get('error') == 0:
                return data.get('data', {}).get('isSubscribe', 0)
            else:
                error_msg = data.get('msg', '未知错误')
                print(f"检查订阅状态失败: {error_msg}")
                
                # 检查是否是登录相关错误
                if '登录' in str(error_msg) or 'login' in str(error_msg).lower():
                    self.handle_login_required_error(error_msg)
                
                return 0
        except requests.RequestException as e:
            print(f"请求订阅状态时发生错误: {e}")
            return 0
    
    def scrape_page(self, res_id, sid):
        """
        抓取完整的页面内容
        :param res_id: 资源ID
        :param sid: 空间ID
        :return: 完整的页面数据
        """
        print(f"开始抓取页面内容: res_id={res_id}, sid={sid}")
        
        # 获取文档内容
        doc_data = self.get_doc_content(res_id, sid)
        if not doc_data:
            print("无法获取文档内容，终止抓取")
            return None
        
        # 获取用户信息
        user_info = self.get_user_info()
        
        # 获取附件
        attachments = self.get_attachments(res_id, sid)
        
        # 检查收藏状态
        collect_status = self.check_collect_status(res_id, sid)
        
        # 检查订阅状态
        subscribe_status = self.check_subscribe_status(res_id)
        
        # 组合完整数据
        result = {
            'doc_data': doc_data,
            'user_info': user_info,
            'attachments': attachments,
            'collect_status': collect_status,
            'subscribe_status': subscribe_status,
            'timestamp': time.time()
        }
        
        print("页面内容抓取完成")
        return result
    
    def save_content_to_file(self, content_data, output_dir="output", formats=["json", "txt", "md", "pdf", "docx"]):
        """
        将抓取的内容保存到文件
        :param content_data: 抓取的数据
        :param output_dir: 输出目录
        :param formats: 要保存的格式列表 (json, txt, md, pdf, docx)
        """
        if not content_data or not content_data.get('doc_data'):
            print("没有有效的内容数据可保存")
            return
        
        doc_data = content_data['doc_data']
        doc_id = doc_data.get('resId', 'unknown')
        
        # 创建输出目录
        os.makedirs(output_dir, exist_ok=True)
        
        # 提取文档内容
        content = doc_data.get('content', '')
        doc_name = doc_data.get('name', f'doc_{doc_id}')
        
        # 保存不同格式的文件
        for fmt in formats:
            if fmt == "json":
                # 保存完整数据为JSON
                json_filename = os.path.join(output_dir, f"doc_{doc_id}_full_data.json")
                with open(json_filename, 'w', encoding='utf-8') as f:
                    json.dump(content_data, f, ensure_ascii=False, indent=2)
                print(f"JSON格式已保存: {json_filename}")
            
            elif fmt == "txt":
                # 保存为文本文件
                txt_filename = os.path.join(output_dir, f"doc_{doc_id}_content.txt")
                with open(txt_filename, 'w', encoding='utf-8') as f:
                    f.write(content)
                print(f"TXT格式已保存: {txt_filename}")
            
            elif fmt == "md":
                # 保存为Markdown文件
                md_filename = os.path.join(output_dir, f"doc_{doc_id}_content.md")
                with open(md_filename, 'w', encoding='utf-8') as f:
                    f.write(content)
                print(f"MD格式已保存: {md_filename}")
            
            elif fmt == "pdf":
                # 保存为PDF文件
                pdf_filename = os.path.join(output_dir, f"doc_{doc_id}_content.pdf")
                self._save_as_pdf(content, pdf_filename, doc_name)
                print(f"PDF格式已保存: {pdf_filename}")
            
            elif fmt == "docx":
                # 保存为DOCX文件
                docx_filename = os.path.join(output_dir, f"doc_{doc_id}_content.docx")
                self._save_as_docx(content, docx_filename, doc_name)
                print(f"DOCX格式已保存: {docx_filename}")
        
        print(f"内容已保存到 {output_dir} 目录，格式: {', '.join(formats)}")
    
    def get_content_summary(self, content_data):
        """
        获取内容摘要信息
        :param content_data: 抓取的数据
        :return: 摘要信息
        """
        if not content_data or not content_data.get('doc_data'):
            return {}
        
        doc_data = content_data['doc_data']
        content = doc_data.get('content', '')
        
        # 统计内容信息
        lines = content.split('\n')
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        # 查找标题
        titles = re.findall(r'^(#+\s+.+)$', content, re.MULTILINE)
        
        summary = {
            'doc_name': doc_data.get('name', ''),
            'creator': doc_data.get('creater', ''),
            'updater': doc_data.get('updater', ''),
            'create_time': doc_data.get('create_time'),
            'update_time': doc_data.get('update_time'),
            'views': doc_data.get('views', 0),
            'content_length': len(content),
            'line_count': len(lines),
            'paragraph_count': len(paragraphs),
            'title_count': len(titles),
            'titles': titles[:10],  # 只返回前10个标题
            'content_preview': content[:200] + '...' if len(content) > 200 else content
        }
        
        return summary
    
    def _save_as_pdf(self, content: str, filename: str, title: str = "文档"):
        """
        将内容保存为PDF格式
        :param content: 文档内容
        :param filename: 输出文件名
        :param title: 文档标题
        """
        if not WEASYPRINT_AVAILABLE:
            print("警告: WeasyPrint不可用，无法生成PDF文件。请安装必要的依赖: pip install weasyprint")
            return
        
        try:
            # 将Markdown转换为HTML
            html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
            
            # 构建完整的HTML文档
            full_html = f'''<!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>{title}</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                    h1, h2, h3, h4, h5, h6 {{ color: #333; }}
                    code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
                    pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
                    table {{ border-collapse: collapse; width: 100%; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                </style>
            </head>
            <body>
                <h1>{title}</h1>
                {html_content}
            </body>
            </html>
            '''
            
            # 使用WeasyPrint转换为PDF
            html_doc = HTML(string=full_html)
            css = CSS(string='''
                @page { margin: 1cm; }
                body { font-size: 12px; }
            ''')
            html_doc.write_pdf(filename, stylesheets=[css])
            
        except Exception as e:
            print(f"保存PDF时发生错误: {e}")
    
    def _save_as_docx(self, content: str, filename: str, title: str = "文档"):
        """
        将内容保存为DOCX格式
        :param content: 文档内容
        :param filename: 输出文件名
        :param title: 文档标题
        """
        try:
            doc = Document()
            
            # 添加标题
            doc.add_heading(title, 0)
            
            # 按行分割内容并添加到文档
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    # 检查是否为标题行
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith('> '):
                        # 引用块
                        doc.add_paragraph(line[2:], style='Intense Quote')
                    elif line.startswith('```'):
                        # 代码块
                        doc.add_paragraph(line, style='Code')
                    else:
                        doc.add_paragraph(line)
            
            doc.save(filename)
            
        except Exception as e:
            print(f"保存DOCX时发生错误: {e}")
    
    def load_cookies_from_config(self):
        """
        从配置文件加载cookie信息
        """
        config_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'api', 'config')
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                config_content = f.read().strip()
                
            # 解析Cookie字符串
            if config_content.startswith('Cookie='):
                cookie_str = config_content[7:]  # 去掉'Cookie='前缀
                
                # 解析cookie键值对
                cookies = {}
                for cookie_pair in cookie_str.split('; '):
                    if '=' in cookie_pair:
                        key, value = cookie_pair.split('=', 1)
                        cookies[key.strip()] = value.strip()
                
                # 将cookie添加到会话中
                self.session.cookies.update(cookies)
                self.login_cookies = cookies
                self.is_logged_in = True
                print("已从配置文件加载登录信息")
            else:
                print("配置文件格式不正确")
                
        except FileNotFoundError:
            print(f"配置文件不存在: {config_path}")
        except Exception as e:
            print(f"加载配置文件时发生错误: {e}")
    
    def login(self, username: str = None, password: str = None) -> bool:
        """
        登录斗鱼文档系统
        :param username: 用户名，如果为None则通过命令行输入
        :param password: 密码，如果为None则通过命令行输入
        :return: 登录是否成功
        """
        if not username:
            username = input("请输入斗鱼文档用户名/邮箱/手机号: ")
        if not password:
            password = getpass.getpass("请输入密码: ")
        
        try:
            # 这里需要根据实际的登录API进行调整
            # 目前根据常见登录模式构建登录请求
            login_url = f"{self.base_url}/login"  # 实际API地址需要根据实际情况调整
            
            # 构建登录数据
            login_data = {
                'username': username,
                'password': password,
                'remember': 1  # 记住登录状态
            }
            
            # 发送登录请求
            response = self.session.post(login_url, data=login_data)
            
            # 检查登录响应
            if response.status_code == 200:
                resp_json = response.json()
                if resp_json.get('error') == 0 or response.url.find('/dashboard') != -1:
                    self.is_logged_in = True
                    self.login_cookies = dict(self.session.cookies)
                    print("登录成功！")
                    return True
                else:
                    print(f"登录失败: {resp_json.get('msg', '未知错误')}")
                    return False
            else:
                print(f"登录请求失败，状态码: {response.status_code}")
                return False
        except Exception as e:
            print(f"登录过程中发生错误: {e}")
            return False
    
    def check_login_status(self) -> bool:
        """
        检查当前登录状态
        :return: 是否已登录
        """
        # 尝试访问需要登录的页面来检查登录状态
        try:
            user_info = self.get_user_info()
            if user_info:
                self.is_logged_in = True
                return True
            else:
                self.is_logged_in = False
                return False
        except:
            self.is_logged_in = False
            return False
    
    def handle_login_required_error(self, error_msg: str):
        """
        处理需要登录的错误
        :param error_msg: 错误消息
        """
        if '登录' in str(error_msg) or 'login' in str(error_msg).lower():
            print("\n检测到需要登录，将使用配置文件中的Cookie重新尝试...")
            # 重新加载配置文件中的Cookie
            self.load_cookies_from_config()


def main():
    # 创建抓取器实例
    scraper = DouyuDocScraper()
    
    # 要抓取的页面参数
    res_id = "251381"
    sid = "539"
    
    print("斗鱼文档内容抓取工具")
    print("="*50)
    
    # 抓取页面内容
    content_data = scraper.scrape_page(res_id, sid)
    
    # 如果抓取失败且可能是登录问题，提示用户登录
    if not content_data:
        print("\n尝试检查登录状态...")
        if not scraper.check_login_status():
            print("检测到未登录状态。")
            scraper.handle_login_required_error("需要登录")
        return
    
    if content_data:
        # 保存内容到文件
        scraper.save_content_to_file(content_data, formats=["json", "txt", "md", "pdf", "docx"])
        
        # 打印文档基本信息
        summary = scraper.get_content_summary(content_data)
        print("\n文档基本信息:")
        print(f"- 文档名称: {summary.get('doc_name', 'N/A')}")
        print(f"- 创建者: {summary.get('creator', 'N/A')}")
        print(f"- 更新者: {summary.get('updater', 'N/A')}")
        print(f"- 创建时间: {summary.get('create_time', 'N/A')}")
        print(f"- 更新时间: {summary.get('update_time', 'N/A')}")
        print(f"- 浏览次数: {summary.get('views', 'N/A')}")
        print(f"- 内容长度: {summary.get('content_length', 0)} 字符")
        print(f"- 行数: {summary.get('line_count', 0)}")
        print(f"- 段落数: {summary.get('paragraph_count', 0)}")
        print(f"- 标题数: {summary.get('title_count', 0)}")
        
        if summary.get('titles'):
            print(f"- 前3个标题: {summary['titles'][:3]}")
        
        print(f"\n内容预览 (前200字符):")
        print(summary.get('content_preview', '')[:200])
    else:
        print("\n抓取失败")
        print("\n可能的原因：")
        print("1. 文档需要登录才能访问")
        print("2. 文档不存在或已被删除")
        print("3. 网络连接问题")
        print("4. API接口发生变化")
        print("\n请检查：")
        print("- 确保您有访问此文档的权限")
        print("- 确保您已登录到斗鱼文档系统")
        print("- 检查网络连接是否正常")


if __name__ == "__main__":
    main()